from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "compliance_validator_code_walkthrough.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
ACCENT = RGBColor(11, 37, 69)
LIGHT_FILL = "F4F6F9"
MID_FILL = "E8EEF5"


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 6),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = BLUE if level <= 2 else DARK_BLUE
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def code_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_fill(cell, LIGHT_FILL)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    doc.add_paragraph()


def two_col_table(doc, rows, headers=("Line / Range", "What it does")):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.55)
    table.columns[1].width = Inches(4.95)
    hdr = table.rows[0].cells
    hdr[0].text = headers[0]
    hdr[1].text = headers[1]
    for line, meaning in rows:
        cells = table.add_row().cells
        cells[0].text = line
        cells[1].text = meaning
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
        if row_idx == 0:
            for cell in row.cells:
                set_cell_fill(cell, MID_FILL)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
    doc.add_paragraph()


def three_col_table(doc, rows, headers=("Lines", "Concept", "Why it matters")):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1.0, 2.2, 3.3]
    for i, w in enumerate(widths):
        table.columns[i].width = Inches(w)
    hdr = table.rows[0].cells
    hdr[0].text = headers[0]
    hdr[1].text = headers[1]
    hdr[2].text = headers[2]
    for row_data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            cells[idx].text = text
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
        if row_idx == 0:
            for cell in row.cells:
                set_cell_fill(cell, MID_FILL)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
    doc.add_paragraph()


def section_header(doc, title, subtitle=None):
    heading(doc, title, 1)
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(subtitle).italic = True


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("AI Compliance Validator Code Walkthrough")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = ACCENT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "A beginner-friendly, line-by-line guide to the code, the concepts, and the build order"
    ).italic = True

    body(
        doc,
        "This document is written for a beginner who wants to understand how the code works, but it explains the project the way an experienced engineer would think about it. The goal is not just to know what each line does, but why that line belongs there.",
    )

    heading(doc, "How To Read This Guide", 1)
    bullets(
        doc,
        [
            "Read the project once from top to bottom before trying to memorize details.",
            "For each file, start with the imports, then the helpers, then the main function.",
            "Think in pipeline order: input, cleaning, retrieval, AI judgment, reporting, UI.",
            "If a line looks small, ask what problem it solves in the larger flow.",
        ],
    )

    heading(doc, "1. The Full Project Idea", 1)
    body(
        doc,
        "The validator is a document-checking pipeline. It takes a policy document and a list of compliance rules, searches the document for evidence, asks an LLM to judge each rule, and returns a structured report.",
    )
    code_box(
        doc,
        "Document + Rules\n"
        "  -> extract text\n"
        "  -> clean text\n"
        "  -> split into chunks\n"
        "  -> store chunks in vector search\n"
        "  -> retrieve relevant evidence\n"
        "  -> ask LLM to judge\n"
        "  -> generate report",
    )

    heading(doc, "2. `main.py` - The CLI Orchestrator", 1)
    body(
        doc,
        "This file is the traffic controller. It does not perform the deep work itself. Instead, it wires together the ingestion layer, the vector store, the validator, and the reporter.",
    )
    two_col_table(
        doc,
        [
            ("1-8", "Module docstring. It tells you the file is the command-line entrypoint and shows example commands."),
            ("10-25", "Imports. These pull in standard library tools, the OpenAI client, the validator modules, and the report helpers."),
            ("28", "Creates a Rich Console object for formatted terminal output."),
            ("31-48", "parse_args() defines the command-line options like `--doc`, `--rules`, `--model`, and `--details`."),
            ("51-52", "main() starts by reading the parsed arguments."),
            ("55-61", "Builds one shared vLLM client and passes it into ingestion and agent modules."),
            ("63-65", "Prints a small banner so the terminal session is readable."),
            ("67-74", "Step 1: load and normalize the document, then chunk it."),
            ("76-79", "Step 2: build the ChromaDB vector store from the chunks."),
            ("81-88", "Step 3: load the rules file."),
            ("90-92", "Step 4: run validation across every rule."),
            ("94-100", "Step 5: build and print the final report, optionally show failure details, and save JSON."),
            ("104-105", "The script runs main() only when you execute the file directly."),
        ],
    )
    body(
        doc,
        "Expert idea: main.py should stay thin. It should read like a recipe, not like a place where every algorithm lives.",
    )

    heading(doc, "3. `validator/ingestion.py` - Read, Clean, Chunk", 1)
    body(
        doc,
        "This file is the document intake system. It converts messy files into clean text and then into chunks that the rest of the project can work with.",
    )
    two_col_table(
        doc,
        [
            ("1-16", "Module docstring. It explains the ingestion pipeline and why AI cleanup is used."),
            ("18-24", "Imports. Regex, JSON, text utilities, Path handling, OpenAI client type, and the LangChain splitter."),
            ("26-35", "Global shared vLLM client state and init_llm(). The client is injected once from main.py."),
            ("40-56", "_extract_pdf_text() reads each PDF page. If no text is found, it raises a helpful OCR message."),
            ("59-65", "_read_text_file() tries several encodings so the project can survive messy files."),
            ("68-74", "_read_html_file() uses BeautifulSoup when available and falls back to regex stripping."),
            ("77-86", "_raw_extract() chooses the correct reader based on file extension."),
            ("91-106", "_fix_encoding() repairs ligatures and removes invisible control characters."),
            ("111-124", "_NORMALISE_SYSTEM is the prompt that tells the LLM how to clean the text."),
            ("127-177", "_ai_normalise() and _normalise_chunk() ask the model to clean layout noise; if that fails, the code keeps going with raw text."),
            ("182-195", "_CHUNK_SYSTEM tells the model how to split the document into meaningful sections."),
            ("198-230", "_ai_chunk() asks the LLM for a JSON list of chunks and falls back to the text splitter if needed."),
            ("233-247", "_fallback_chunk() uses RecursiveCharacterTextSplitter for safer chunking when AI chunking is unavailable."),
            ("252-261", "load_document() runs raw extraction -> encoding fix -> AI normalization."),
            ("264-275", "load_and_chunk() runs the same pipeline but ends with chunk creation instead of plain text."),
            ("278-304", "load_rules() reads the rules file and chooses AI extraction or regex fallback."),
            ("307-325", "_RULES_SYSTEM is the prompt for converting messy rule text into structured JSON."),
            ("328-361", "_ai_extract_rules() asks the LLM to parse rules into id/text dictionaries and falls back safely."),
            ("364-389", "_regex_extract_rules() is the backup parser when the LLM is unavailable."),
            ("393-395", "chunk_text() is a backward-compatible alias that keeps older code working."),
        ],
    )
    body(
        doc,
        "Expert idea: this file is layered on purpose. Always do the simplest deterministic cleanup first, then use AI only where it adds real value.",
    )

    heading(doc, "4. `validator/vector_store.py` - Semantic Search", 1)
    body(
        doc,
        "This file turns chunks into searchable memory. The goal is to find the most relevant pieces of the document for a given rule.",
    )
    two_col_table(
        doc,
        [
            ("1-6", "Defines the embedding model and creates the embedding function."),
            ("9-26", "build_store() creates an in-memory ChromaDB collection and inserts all chunks."),
            ("29-39", "retrieve() searches the collection with a rule query and returns the top matching chunks."),
        ],
    )
    body(
        doc,
        "Expert idea: semantic search is the bridge between the checklist and the document. It keeps the LLM from reading everything blindly.",
    )

    heading(doc, "5. `validator/agent.py` - The LLM Judge", 1)
    body(
        doc,
        "This file is the decision-maker. It asks the model to judge whether a document satisfies one rule at a time.",
    )
    two_col_table(
        doc,
        [
            ("1-17", "Creates the shared OpenAI client and default model name."),
            ("23-42", "SYSTEM_PROMPT defines the exact JSON format and the meaning of PASS, FAIL, and UNCLEAR."),
            ("45-63", "_parse_llm_response() strips markdown fences and safely parses JSON, with a fallback if parsing fails."),
            ("66-77", "validate_rule() describes the job: retrieve chunks, send them to the model, and return a structured result."),
            ("79-87", "The retrieved chunks are joined into a short context block for the model."),
            ("89-97", "The chat completion call asks the LLM for a low-temperature, structured answer."),
            ("99-110", "The parsed response is turned into a uniform result dictionary with supporting evidence."),
            ("113-132", "run_validation() loops through all rules, prints progress, and gathers results."),
        ],
    )
    code_box(
        doc,
        "Rule -> retrieve evidence -> ask LLM -> parse JSON -> store result",
    )
    body(
        doc,
        "Expert idea: the prompt is doing a lot of work here. Good prompts are like precise instructions for a careful assistant, not vague requests.",
    )

    heading(doc, "6. `validator/reporter.py` - Summary and Storytelling", 1)
    body(
        doc,
        "This file turns raw findings into a report a human can understand quickly. It also adds risk labels and recommendations, which are important for demos and practical use.",
    )
    two_col_table(
        doc,
        [
            ("13-19", "_risk_label() converts compliance score into Low, Medium, or High Risk."),
            ("22-27", "_risk_color() maps the risk label to a display color."),
            ("30-41", "_recommendation_for_finding() creates simple remediation advice."),
            ("44-53", "_sort_findings() moves FAIL and UNCLEAR before PASS."),
            ("56-91", "generate_report() computes totals, score, risk label, and adds recommendations to every finding."),
            ("94-138", "print_terminal_report() renders the summary panel and the findings table."),
            ("141-160", "print_failed_details() prints a deeper explanation for every failed rule."),
            ("163-169", "save_report() writes the final JSON report to disk."),
        ],
    )
    body(
        doc,
        "Expert idea: a report is not just output. It is the interface between the system and the human making a decision.",
    )

    heading(doc, "7. `app.py` - The Streamlit UI", 1)
    body(
        doc,
        "This file is the visual wrapper around the same backend pipeline. It lets a user upload files and inspect the report in a browser.",
    )
    three_col_table(
        doc,
        [
            ("1-17", "Imports and setup", "Connects Streamlit with ingestion, vector search, validation, and reporting."),
            ("18-26", "Page config", "Sets the title, icon, and wide layout."),
            ("28-40", "Sidebar settings", "Lets the user change the model, vLLM URL, top-k, and chunk size."),
            ("42-49", "File uploads", "Accepts the document and rules file."),
            ("51-81", "Run validation flow", "Writes temp files, loads text, chunks, builds store, loads rules, validates, and creates the report."),
            ("83-92", "Summary metrics", "Shows compliance score, risk level, and counts."),
            ("94-107", "Findings display", "Shows each rule in an expander with evidence and recommendation."),
            ("109-115", "Download button", "Lets the user save the JSON report."),
        ],
    )
    body(
        doc,
        "Expert idea: the UI should not hide the core logic. It should simply make the same pipeline easier to use.",
    )

    heading(doc, "8. The Build Order An Expert Would Use", 1)
    numbered(
        doc,
        [
            "Start with sample document and sample rules.",
            "Make extraction work first.",
            "Add cleaning and chunking.",
            "Add vector search.",
            "Make one rule pass end to end.",
            "Add the LLM judge.",
            "Loop over all rules.",
            "Generate the report.",
            "Add terminal output.",
            "Add the UI only after the backend is stable.",
        ],
    )

    heading(doc, "9. How To Debug This Project", 1)
    two_col_table(
        doc,
        [
            ("No text extracted", "Check the file type, encoding, or whether the PDF is scanned."),
            ("Wrong chunks retrieved", "Print the retrieved chunks and inspect whether the semantic search is relevant."),
            ("Invalid JSON from LLM", "Tighten the prompt and keep the JSON fallback parser."),
            ("Everything is UNCLEAR", "The rule may be vague or the retrieved evidence may be too weak."),
            ("Slow demo", "Lower top-k or reduce max_tokens."),
        ],
        headers=("Problem", "First thing to check"),
    )

    heading(doc, "10. Reusable Pattern For Future Projects", 1)
    body(
        doc,
        "This is the general shape you can reuse for many AI tools: take documents or records, break them into pieces, retrieve the relevant pieces, ask AI a focused question, and save structured output.",
    )
    code_box(
        doc,
        "1. Collect input\n"
        "2. Clean input\n"
        "3. Split into chunks\n"
        "4. Store chunks for retrieval\n"
        "5. Retrieve relevant evidence\n"
        "6. Ask AI a focused question\n"
        "7. Force structured output\n"
        "8. Parse and validate output\n"
        "9. Produce a report\n"
        "10. Add the UI last",
    )

    heading(doc, "11. Beginner Checklist", 1)
    bullets(
        doc,
        [
            "Can I explain what each file is responsible for?",
            "Can I trace one document from upload to final report?",
            "Can I describe why RAG is used instead of sending the whole document?",
            "Can I explain why structured JSON matters?",
            "Can I point to where the risk label and recommendation are created?",
        ],
    )

    body(
        doc,
        "If you read this guide alongside the code, you should be able to answer not only what the project does, but how each part fits into the whole. That is the point where the code starts to feel learnable instead of mysterious.",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
