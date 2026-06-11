from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "compliance_validator_beginner_guide.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def format_table(table, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
        if header and row_idx == 0:
            for cell in row.cells:
                set_cell_fill(cell, GRAY_FILL)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = BLUE if level <= 2 else DARK_BLUE
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_fill(cell, LIGHT_FILL)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    p.add_run("\n" + body)
    doc.add_paragraph()


def add_two_col_table(doc, rows, widths=(2.0, 4.5), headers=("Part", "Meaning")):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Inches(widths[0])
    table.columns[1].width = Inches(widths[1])
    table.cell(0, 0).text = headers[0]
    table.cell(0, 1).text = headers[1]
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
    format_table(table)
    doc.add_paragraph()


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_fill(cell, "F7F7F7")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    doc.add_paragraph()


def set_document_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_document_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI Compliance Validator: Beginner Build Guide")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = DARK_BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "A practical, expert-style walkthrough for understanding and building similar AI projects"
    ).italic = True

    add_callout(
        doc,
        "Purpose of this guide",
        "This guide explains the project from a beginner's point of view, but with the thinking process of an experienced builder. Read it as a map: what to build first, why each step exists, and how to reuse the pattern in future projects.",
    )

    add_heading(doc, "1. The Project In One Simple Idea")
    add_body(
        doc,
        "AI Compliance Validator is an AI-powered checklist inspector. You give it a document and a list of compliance rules. It checks whether the document satisfies each rule, then produces a report with pass/fail/unclear verdicts, evidence, confidence, risk level, and recommendations.",
    )
    add_code_block(
        doc,
        "Document + Rules\n"
        "    -> read and clean text\n"
        "    -> split document into chunks\n"
        "    -> search relevant chunks for each rule\n"
        "    -> ask AI to judge using evidence\n"
        "    -> generate audit report",
    )

    add_heading(doc, "2. How An Expert Thinks Before Coding")
    add_body(
        doc,
        "An expert does not start by opening Streamlit or writing prompts. They first define the data flow. They ask: What is the input? What is the output? What steps turn the input into the output?",
    )
    add_two_col_table(
        doc,
        [
            ("Input", "A policy document and a compliance rules file."),
            ("Output", "A structured audit report that a human can understand and trust."),
            ("Core job", "Check every rule against evidence from the document."),
            ("AI job", "Understand language, judge compliance, explain the decision."),
            ("Normal code job", "Read files, clean text, search chunks, calculate scores, display reports."),
        ],
        headers=("Question", "Answer"),
    )
    add_callout(
        doc,
        "Expert habit",
        "Separate normal programming from AI reasoning. Use code for deterministic work. Use AI when the task needs language understanding.",
    )

    add_heading(doc, "3. The End-To-End Pipeline")
    add_numbered(
        doc,
        [
            "Read the document from PDF, TXT, Markdown-like text, or HTML.",
            "Clean the extracted text so the AI sees readable content.",
            "Split the document into smaller chunks.",
            "Store chunks in ChromaDB so they can be searched by meaning.",
            "Read the compliance rules from a text file.",
            "For each rule, retrieve the most relevant document chunks.",
            "Send the rule and evidence chunks to the LLM judge.",
            "Parse the LLM's JSON response.",
            "Calculate summary metrics, risk label, and recommendations.",
            "Show results in the CLI or Streamlit UI.",
        ],
    )

    add_heading(doc, "4. Why The Project Uses RAG")
    add_body(
        doc,
        "RAG means Retrieval Augmented Generation. In beginner language: first find the relevant parts of the document, then ask the AI a focused question using only those parts.",
    )
    add_two_col_table(
        doc,
        [
            ("Without RAG", "The AI may receive too much text, miss details, or guess."),
            ("With RAG", "The AI receives the rule plus the most relevant evidence."),
            ("Benefit", "Better accuracy, lower token usage, and clearer audit evidence."),
        ],
        headers=("Approach", "Meaning"),
    )

    add_heading(doc, "5. File-By-File Mental Model")
    add_two_col_table(
        doc,
        [
            ("main.py", "Command-line entrypoint. Runs the full pipeline from terminal."),
            ("app.py", "Streamlit UI. Lets users upload files and view results visually."),
            ("validator/ingestion.py", "Reads documents, cleans text, chunks text, and loads rules."),
            ("validator/vector_store.py", "Builds the ChromaDB smart-search index and retrieves chunks."),
            ("validator/agent.py", "Talks to the LLM and asks it to judge each rule."),
            ("validator/reporter.py", "Creates the final report, risk label, sorted findings, and recommendations."),
            ("rules/compliance_rules.txt", "Default checklist of compliance requirements."),
            ("documents/*.txt", "Sample documents used for demos and testing."),
        ],
        headers=("File", "Responsibility"),
    )

    add_heading(doc, "6. Step 1: Start With Sample Data")
    add_body(
        doc,
        "A good builder creates tiny test data before building the whole system. This project has a compliant policy and a non-compliant policy. This makes it easy to check whether the system is working.",
    )
    add_bullets(
        doc,
        [
            "documents/sample_policy.txt should mostly pass.",
            "documents/non_compliant_sample.txt should fail many rules.",
            "rules/compliance_rules.txt defines what the system checks.",
        ],
    )
    add_callout(
        doc,
        "Why this matters",
        "You cannot improve a system if you cannot test it. Sample data gives you a quick reality check after every change.",
    )

    add_heading(doc, "7. Step 2: Read And Clean The Document")
    add_body(
        doc,
        "The ingestion layer converts files into plain text. AI models do not understand a PDF file path. They need readable text. That is why the first real step is extraction.",
    )
    add_two_col_table(
        doc,
        [
            ("PDF", "Use pypdf to extract embedded text."),
            ("TXT/MD", "Read the text directly with common encodings."),
            ("HTML", "Use BeautifulSoup to remove tags and keep visible text."),
            ("Scanned PDF", "Needs OCR first because there is no embedded text to extract."),
        ],
        headers=("File Type", "What Happens"),
    )
    add_body(
        doc,
        "After extraction, the project cleans broken characters, odd spaces, and layout noise. This is important because messy input produces unreliable AI output.",
    )

    add_heading(doc, "8. Step 3: Split Text Into Chunks")
    add_body(
        doc,
        "Large documents are hard for AI to process at once. So the document is split into chunks, such as coverage details, exclusions, renewal terms, or cancellation terms.",
    )
    add_code_block(
        doc,
        "Policy document\n"
        "    -> Chunk 1: Policyholder information\n"
        "    -> Chunk 2: Coverage details\n"
        "    -> Chunk 3: Premium terms\n"
        "    -> Chunk 4: Exclusions\n"
        "    -> Chunk 5: Cancellation terms",
    )
    add_body(
        doc,
        "The goal is not just smaller text. The goal is meaningful smaller text. A rule about exclusions should retrieve the exclusions chunk, not random policyholder information.",
    )

    add_heading(doc, "9. Step 4: Build Smart Search With ChromaDB")
    add_body(
        doc,
        "ChromaDB stores chunks as embeddings. An embedding is a numerical representation of meaning. This lets the project search by meaning, not only exact words.",
    )
    add_two_col_table(
        doc,
        [
            ("Exact search", "Looks for the same words."),
            ("Vector search", "Looks for similar meaning."),
            ("Example", "A rule mentioning 'exclusions' can find text saying 'not covered under this policy'."),
        ],
        headers=("Concept", "Explanation"),
    )

    add_heading(doc, "10. Step 5: Load Rules As Data")
    add_body(
        doc,
        "The rules are not hardcoded in Python. They live in a separate rules file. This is good design because users can change the checklist without changing the application code.",
    )
    add_code_block(
        doc,
        "RULE001: Document must explicitly state the coverage start and end date\n"
        "RULE002: Policy must include the full legal name of the insured party\n"
        "RULE003: Document must disclose all exclusion clauses clearly",
    )
    add_body(
        doc,
        "The code converts these lines into dictionaries with an id and text. That structure makes it easy to loop over rules one by one.",
    )

    add_heading(doc, "11. Step 6: Judge One Rule At A Time")
    add_body(
        doc,
        "The agent layer takes one rule, retrieves relevant chunks, and asks the LLM for a structured decision. The LLM is instructed to act like a strict compliance auditor.",
    )
    add_code_block(
        doc,
        "Rule:\n"
        "Policy must include the full legal name of the insured party.\n\n"
        "Relevant evidence:\n"
        "Full Legal Name: John Michael Doe\n\n"
        "Expected AI output:\n"
        "{\n"
        '  \"verdict\": \"PASS\",\n'
        '  \"confidence\": 0.97,\n'
        '  \"evidence\": \"Full Legal Name: John Michael Doe\",\n'
        '  \"explanation\": \"The document includes the full legal name.\"\n'
        "}",
    )
    add_callout(
        doc,
        "Important AI engineering principle",
        "Ask the model for JSON, not a casual paragraph. Structured output is easier for code to parse, display, save, and test.",
    )

    add_heading(doc, "12. Step 7: Generate A Human-Friendly Report")
    add_body(
        doc,
        "The reporter layer turns individual rule decisions into a full audit report. It calculates totals, compliance score, average confidence, risk level, and recommendations.",
    )
    add_two_col_table(
        doc,
        [
            ("Compliance score", "Passed rules divided by total rules."),
            ("Risk label", "Low Risk, Medium Risk, or High Risk based on score."),
            ("Findings order", "FAIL first, UNCLEAR second, PASS last."),
            ("Recommendation", "A short action telling the user how to fix failed or unclear rules."),
        ],
        headers=("Report Field", "Why It Exists"),
    )

    add_heading(doc, "13. Step 8: Add Interfaces")
    add_body(
        doc,
        "The project has two interfaces. The CLI is for fast developer testing. The Streamlit UI is for demos and non-technical users.",
    )
    add_two_col_table(
        doc,
        [
            ("CLI", "Best for quick testing, automation, and debugging."),
            ("Streamlit", "Best for hackathon judges and visual demos."),
        ],
        headers=("Interface", "Use Case"),
    )
    add_body(
        doc,
        "An expert usually builds the backend pipeline first, then adds the UI. This prevents the UI from hiding broken logic.",
    )

    add_heading(doc, "14. The Expert Build Order")
    add_numbered(
        doc,
        [
            "Create one sample document and one rules file.",
            "Write code to read the document into text.",
            "Write code to read rules into a list.",
            "Split the document into chunks.",
            "Build vector search over chunks.",
            "Retrieve chunks for one rule.",
            "Ask the LLM to judge one rule.",
            "Parse the LLM response as JSON.",
            "Loop over all rules.",
            "Generate a JSON report.",
            "Add terminal output.",
            "Add Streamlit UI.",
            "Add polish: risk labels, recommendations, sorting, README, demo script.",
        ],
    )

    add_heading(doc, "15. How To Debug This Kind Of Project")
    add_body(
        doc,
        "When something fails, debug the pipeline step by step. Do not guess at the end. Check each transformation.",
    )
    add_two_col_table(
        doc,
        [
            ("Bad or empty report", "Check whether the document text was extracted correctly."),
            ("Wrong chunks retrieved", "Print retrieved chunks and inspect whether ChromaDB found useful evidence."),
            ("LLM gives invalid JSON", "Tighten the prompt and add JSON parsing fallback."),
            ("Everything is UNCLEAR", "Check whether rules are too vague or evidence chunks are missing."),
            ("Slow demo", "Lower top-k or reduce max_tokens in the LLM call."),
        ],
        headers=("Problem", "First Place To Look"),
    )

    add_heading(doc, "16. Reusable Pattern For Future AI Projects")
    add_body(
        doc,
        "This architecture is useful far beyond compliance validation. You can reuse it whenever you need AI to inspect a large document using a checklist or question set.",
    )
    add_code_block(
        doc,
        "1. Collect input\n"
        "2. Clean input\n"
        "3. Split into chunks\n"
        "4. Store chunks for semantic search\n"
        "5. Retrieve relevant context\n"
        "6. Ask AI a focused question\n"
        "7. Force structured JSON output\n"
        "8. Parse and validate output\n"
        "9. Generate a human-friendly report\n"
        "10. Add a UI only after the backend works",
    )
    add_bullets(
        doc,
        [
            "Resume evaluator: resume + job requirements -> fit report.",
            "Contract checker: contract + legal checklist -> risk report.",
            "Invoice validator: invoice + policy rules -> exception report.",
            "Research assistant: papers + questions -> cited answers.",
            "Support bot: knowledge base + user question -> grounded response.",
        ],
    )

    add_heading(doc, "17. Beginner Checklist Before Building A Similar Project")
    add_bullets(
        doc,
        [
            "Can I explain the input and output in one sentence?",
            "Do I have sample data that should pass and sample data that should fail?",
            "Can I run the backend without a UI?",
            "Can I print each intermediate step for debugging?",
            "Am I using normal code for deterministic tasks?",
            "Am I using AI only where language understanding is needed?",
            "Is the AI response structured enough for code to parse?",
            "Does the final report help a human take action?",
        ],
    )

    add_heading(doc, "18. Final Mental Model")
    add_callout(
        doc,
        "One-line summary",
        "First make the data flow work end to end, then make each step smarter, then make the output easy for humans to trust.",
    )
    add_body(
        doc,
        "If you remember only one thing, remember this: build AI projects as pipelines. Each step should have a clear input, a clear output, and a reason to exist.",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
